from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_project_report():
    document = Document()

    # Title
    title = document.add_heading('Image Clustering & Retrieval System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph('Project Report & Documentation', style='Subtitle')
    document.add_paragraph('Developed for: University Assignment', style='Body Text')

    # 1. Introduction
    document.add_heading('1. Introduction', level=1)
    p = document.add_paragraph()
    p.add_run('This project is an ').bold = True
    p.add_run('AI-powered application').bold = True
    p.add_run(' designed to automatically organize, classify, and search through a large collection of images. Instead of manually sorting thousands of photos, this system uses "Computer Vision" to understand the content of images and group them based on visual similarity.')

    # 2. What is this Project Doing? (The Core Logic)
    document.add_heading('2. How It Works (The Core Logic)', level=1)
    document.add_paragraph("The system follows a standard Machine Learning pipeline:")

    # 2.1 Feature Extraction
    document.add_heading('2.1. Feature Extraction (The "Eyes")', level=2)
    document.add_paragraph('Computers cannot "see" images like humans do; they only see grids of pixels. To solve this, we use a technique called Feature Extraction.')
    document.add_paragraph('• We use a pre-trained Deep Learning model called ResNet18 (a Convolutional Neural Network).')
    document.add_paragraph('• This model converts every image into a list of 512 numbers (a "vector").')
    document.add_paragraph('• These numbers represent the "essence" of the image (e.g., shapes, textures, colors) rather than just raw pixels.')

    # 2.2 Clustering
    document.add_heading('2.2. Clustering (Unsupervised Learning)', level=2)
    document.add_paragraph('Once every image is converted into numbers, we can group them.')
    document.add_paragraph('• K-Means Clustering: We tell the computer "find 5 groups," and it mathematically finds the best way to separate the images into 5 distinct piles based on similarity.')
    document.add_paragraph('• Hierarchical Clustering: This builds a "family tree" of images, merging similar ones step-by-step.')
    document.add_paragraph('• Purpose: To organize data without needing any human labels.')

    # 2.3 Evaluation
    document.add_heading('2.3. Evaluation (Grading the AI)', level=2)
    document.add_paragraph('How do we know if the groups are good? We use mathematical scores:')
    document.add_paragraph('• Silhouette Score: Measures how similar an image is to its own cluster compared to other clusters. (Higher is better, close to 1).')
    document.add_paragraph('• Davies-Bouldin Index: Measures how "separated" the clusters are. (Lower is better).')

    # 2.4 Classification
    document.add_heading('2.4. Classification (Supervised Learning)', level=2)
    document.add_paragraph('If we have some labeled examples (e.g., "this is a red circle"), we can train a model to predict labels for new images.')
    document.add_paragraph('• We use SVM (Support Vector Machine) and Random Forest algorithms.')
    document.add_paragraph('• This allows the system to say "This new image is likely a Green Triangle".')

    # 2.5 Retrieval
    document.add_heading('2.5. Image Retrieval (Search Engine)', level=2)
    document.add_paragraph('This is like a "Google Image Search" for your dataset.')
    document.add_paragraph('• You upload a photo.')
    document.add_paragraph('• The system calculates its feature vector.')
    document.add_paragraph('• It looks for the "nearest neighbors" (mathematically closest vectors) in the database.')
    document.add_paragraph('• It returns the most similar images found.')

    # 3. Project Structure
    document.add_heading('3. Code Structure', level=1)
    table = document.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'File Name'
    hdr_cells[1].text = 'Purpose'

    files = [
        ('app.py', 'The main website/interface (built with Streamlit). This is what you run to see the project.'),
        ('feature_extraction.py', 'Contains the AI brain (ResNet18) to read images.'),
        ('clustering.py', 'Contains the math for K-Means and Hierarchical clustering.'),
        ('evaluation.py', 'Calculates the scores (Silhouette, etc.) to grade performance.'),
        ('classifier.py', 'Contains the SVM and Random Forest logic for prediction.'),
        ('retrieval.py', 'The search engine logic to find similar images.'),
        ('generate_synthetic_data.py', 'A helper script to create dummy images (shapes/colors) for testing.')
    ]

    for name, desc in files:
        row_cells = table.add_row().cells
        row_cells[0].text = name
        row_cells[1].text = desc

    # 4. Conclusion
    document.add_heading('4. Conclusion', level=1)
    document.add_paragraph('This project demonstrates a complete end-to-end Computer Vision pipeline. It successfully integrates Deep Learning (for features) with Classical Machine Learning (for clustering/classification) to solve the problem of organizing and searching large image datasets.')

    document.save('Project_Report.docx')
    print("Document created successfully.")

if __name__ == "__main__":
    create_project_report()
