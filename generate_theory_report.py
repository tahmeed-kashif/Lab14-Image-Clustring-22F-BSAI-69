from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_theory_report():
    document = Document()

    # Title
    title = document.add_heading('Machine Learning Algorithms Explained', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    document.add_paragraph('Detailed explanation of the algorithms used in the Image Clustering & Retrieval Project.', style='Subtitle')

    # 1. K-Means Clustering
    document.add_heading('1. K-Means Clustering', level=1)
    document.add_paragraph('Purpose:', style='Heading 3')
    document.add_paragraph('K-Means is an unsupervised learning algorithm used to group similar data points (images) together without knowing their labels beforehand. In this project, it organizes images into distinct groups (e.g., separating apples from cats).')
    
    document.add_paragraph('How it Works:', style='Heading 3')
    p = document.add_paragraph()
    p.add_run('1. Initialization: ').bold = True
    p.add_run('The algorithm randomly selects "K" center points (centroids).')
    
    p = document.add_paragraph()
    p.add_run('2. Assignment: ').bold = True
    p.add_run('Every image is compared to these centroids. The image is assigned to the cluster whose centroid is closest (using distance formulas).')
    
    p = document.add_paragraph()
    p.add_run('3. Update: ').bold = True
    p.add_run('Once all images are assigned, the algorithm calculates the average position of all images in a cluster and moves the centroid to this new center.')
    
    p = document.add_paragraph()
    p.add_run('4. Repeat: ').bold = True
    p.add_run('Steps 2 and 3 are repeated until the centroids stop moving. At this point, the clusters are stable.')

    # 2. Hierarchical Clustering
    document.add_heading('2. Hierarchical Clustering', level=1)
    document.add_paragraph('Purpose:', style='Heading 3')
    document.add_paragraph('This algorithm builds a hierarchy of clusters. Unlike K-Means, you don’t strictly need to pick the number of clusters upfront; it builds a tree-like structure (dendrogram) showing how images are related at different levels of similarity.')
    
    document.add_paragraph('How it Works (Agglomerative Approach):', style='Heading 3')
    p = document.add_paragraph()
    p.add_run('1. Start: ').bold = True
    p.add_run('Treat every single image as its own tiny cluster.')
    
    p = document.add_paragraph()
    p.add_run('2. Merge: ').bold = True
    p.add_run('Find the two most similar clusters (closest together) and merge them into one larger cluster.')
    
    p = document.add_paragraph()
    p.add_run('3. Repeat: ').bold = True
    p.add_run('Keep merging the closest pairs until all images are combined into one giant cluster.')
    
    document.add_paragraph('The result is a tree structure. To get specific groups (like 5 clusters), we simply "cut" the tree at a specific height.')

    # 3. Support Vector Machine (SVM)
    document.add_heading('3. Support Vector Machine (SVM)', level=1)
    document.add_paragraph('Purpose:', style='Heading 3')
    document.add_paragraph('SVM is a supervised learning algorithm used for classification. In this project, it learns to distinguish between labeled categories (e.g., "Apple" vs "Banana") so it can predict the label of a new, unseen image.')
    
    document.add_paragraph('How it Works:', style='Heading 3')
    document.add_paragraph('Imagine plotting the image data points on a graph. SVM tries to find the best line (or "hyperplane" in 3D+) that separates the two classes.')
    
    p = document.add_paragraph()
    p.add_run('• The Margin: ').bold = True
    p.add_run('SVM looks for the widest possible "street" (margin) between the classes. It wants the line that is furthest away from the closest dots of both colors.')
    
    p = document.add_paragraph()
    p.add_run('• Support Vectors: ').bold = True
    p.add_run('The data points closest to the line are called "support vectors" because they support or define where the line goes.')
    
    document.add_paragraph('If the data isn’t easily separable with a straight line, SVM uses a "Kernel Trick" to project data into higher dimensions where it becomes separable.')

    # 4. Random Forest
    document.add_heading('4. Random Forest', level=1)
    document.add_paragraph('Purpose:', style='Heading 3')
    document.add_paragraph('Random Forest is another supervised classifier. It is an "ensemble" method, meaning it combines many small models to make a better decision. It is generally very robust and accurate.')
    
    document.add_paragraph('How it Works:', style='Heading 3')
    p = document.add_paragraph()
    p.add_run('1. Decision Trees: ').bold = True
    p.add_run('It builds many "Decision Trees". A decision tree asks a series of Yes/No questions about the image features (e.g., "Is the center pixel red?") to arrive at a label.')
    
    p = document.add_paragraph()
    p.add_run('2. Randomness: ').bold = True
    p.add_run('Each tree is trained on a random subset of the data and looks at a random subset of features. This ensures the trees are different from each other.')
    
    p = document.add_paragraph()
    p.add_run('3. Voting: ').bold = True
    p.add_run('When we want to classify a new image, every tree in the forest makes a prediction. The Random Forest counts the votes and chooses the label with the most votes (Majority Voting).')

    document.save('Algorithm_Explanation_Report.docx')
    print("Theory report created successfully.")

if __name__ == "__main__":
    create_theory_report()
